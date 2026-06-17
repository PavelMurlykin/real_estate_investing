from collections import OrderedDict
from copy import deepcopy
from dataclasses import dataclass
from datetime import date, datetime
from io import BytesIO
from pathlib import Path

from django.http import HttpResponse
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RELATIONSHIP_TYPE
from docx.shared import Inches, Pt, RGBColor

from .excel import (
    _build_property_rows,
    build_saved_mortgage_excel_data,
    get_discount_markup_labels,
)
from .utils import format_compact_decimal, format_currency


WORD_CONTENT_TYPE = (
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
)
WORD_TEMPLATE_PATH = (
    Path(__file__).resolve().parent
    / 'word_templates'
    / 'mortgage_report_template.docx'
)
COMPLEX_ROW_LABELS = {
    'Город',
    'Район',
    'Застройщик',
    'Жилой комплекс',
    'Название ЖК',
    'Класс ЖК',
    'Сдача',
    'Ключи',
}


@dataclass
class MortgageWordCalculation:
    """Prepared calculation data for Word mortgage reports."""

    property_obj: object
    template_rows: dict
    calculation_heading: str
    program_label: str


def export_mortgage_word(report_data):
    """Формирует HTTP-ответ с Word-файлом ипотечного расчета."""
    document = _build_market_mortgage_document(report_data)
    return _build_word_response(document, 'mortgage_calculation.docx')


def export_saved_mortgage_calculation_word(calculation, payment_schedule):
    """Формирует Word-файл для сохраненного ипотечного расчета."""
    report_data = build_saved_mortgage_excel_data(
        calculation,
        payment_schedule,
    )
    return export_mortgage_word(report_data)


def export_trench_mortgage_word(calculation):
    """Формирует HTTP-ответ с Word-файлом траншевого расчета."""
    document = _build_trench_mortgage_document(calculation)
    return _build_word_response(
        document,
        'trench_mortgage_calculation.docx',
    )


def export_customer_mortgage_calculations_word(calculations):
    """Build a grouped Word report for selected customer calculations."""
    prepared_calculations = list(calculations)
    document = _build_customer_mortgage_document(prepared_calculations)
    return _build_word_response(
        document,
        'customer_mortgage_calculations.docx',
    )


def build_market_word_calculation(calculation, payment_schedule):
    """Build prepared Word data for a saved market mortgage calculation."""
    report_data = build_saved_mortgage_excel_data(
        calculation,
        payment_schedule,
    )
    return MortgageWordCalculation(
        property_obj=report_data.property_obj,
        template_rows=_build_market_template_rows(report_data),
        calculation_heading='Ипотека',
        program_label='Рыночная ипотека',
    )


def build_trench_word_calculation(calculation):
    """Build prepared Word data for a saved trench mortgage calculation."""
    return MortgageWordCalculation(
        property_obj=calculation.get('property_obj'),
        template_rows=_build_trench_template_rows(calculation),
        calculation_heading='Ипотека траншевая',
        program_label='Траншевая ипотека',
    )


def _build_word_response(document, filename):
    """Сохраняет Word-документ в HTTP-ответ для скачивания."""
    output = BytesIO()
    document.save(output)
    response = HttpResponse(output.getvalue(), content_type=WORD_CONTENT_TYPE)
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


def _build_market_mortgage_document(report_data):
    """Создает Word-документ для обычного ипотечного расчета."""
    document = _create_template_document()
    _populate_template_document(
        document,
        report_data.property_obj,
        _build_market_template_rows(report_data),
        'Ипотека',
    )
    return document


def _build_trench_mortgage_document(calculation):
    """Создает Word-документ для траншевого ипотечного расчета."""
    document = _create_template_document()
    _populate_template_document(
        document,
        calculation.get('property_obj'),
        _build_trench_template_rows(calculation),
        'Ипотека траншевая',
    )
    return document


def _build_customer_mortgage_document(calculations):
    """Create a grouped Word document for customer selected calculations."""
    document = _create_template_document()
    if not calculations:
        _clear_document_body(document)
        return document

    if len(document.tables) < 3:
        _populate_customer_fallback_document(document, calculations)
        return document

    template_tables = [
        deepcopy(table._tbl)
        for table in document.tables[:3]
    ]
    grouped_calculations = _group_word_calculations(calculations)
    _clear_document_body(document)

    is_first_complex = True
    for complex_group in grouped_calculations:
        if not is_first_complex:
            _append_page_break(document)
        is_first_complex = False

        first_property_group = complex_group['property_groups'][0]
        complex_table = _append_template_table(
            document,
            template_tables[0],
        )
        _populate_complex_table(
            complex_table,
            first_property_group['property_obj'],
            first_property_group['template_rows']['complex_rows'],
        )

        for property_index, property_group in enumerate(
            complex_group['property_groups']
        ):
            if property_index > 0:
                _append_page_break(document)
            _append_customer_property_tables(
                document,
                template_tables,
                property_group,
            )

    return document


def _append_customer_property_tables(document, template_tables, property_group):
    """Append and fill object, calculation and image tables."""
    object_table = _append_template_table(document, template_tables[1])
    image_table = _append_template_table(document, template_tables[2])
    _populate_object_and_calculation_table(
        object_table,
        property_group['property_obj'],
        property_group['template_rows'],
        property_group['calculation_heading'],
    )
    _populate_image_table(image_table, property_group['property_obj'])


def _populate_customer_fallback_document(document, calculations):
    """Populate a simple customer report when the Word template is invalid."""
    _clear_document_body(document)
    grouped_calculations = _group_word_calculations(calculations)
    is_first_section = True
    for complex_group in grouped_calculations:
        if not is_first_section:
            _append_page_break(document)
        is_first_section = False
        first_property_group = complex_group['property_groups'][0]
        _append_key_value_section(
            document,
            'Данные ЖК',
            first_property_group['template_rows']['complex_rows'],
        )
        for property_index, property_group in enumerate(
            complex_group['property_groups']
        ):
            if property_index > 0:
                _append_page_break(document)
            _append_key_value_section(
                document,
                'Объект недвижимости',
                property_group['template_rows']['object_rows'],
            )
            _append_key_value_section(
                document,
                property_group['calculation_heading'],
                property_group['template_rows']['calculation_rows'],
            )


def _group_word_calculations(calculations):
    """Group calculations by complex and property preserving selection order."""
    complex_groups = OrderedDict()
    for calculation in calculations:
        complex_key = _get_word_complex_key(calculation.property_obj)
        property_key = _get_word_property_key(calculation.property_obj)

        if complex_key not in complex_groups:
            complex_groups[complex_key] = OrderedDict()
        property_groups = complex_groups[complex_key]
        if property_key not in property_groups:
            property_groups[property_key] = {
                'property_obj': calculation.property_obj,
                'calculations': [],
            }
        property_groups[property_key]['calculations'].append(calculation)

    return [
        {
            'property_groups': [
                _build_word_property_group(property_group)
                for property_group in property_groups.values()
            ]
        }
        for property_groups in complex_groups.values()
    ]


def _build_word_property_group(property_group):
    """Build merged template rows for one property calculation group."""
    calculations = property_group['calculations']
    first_calculation = calculations[0]
    template_rows = {
        'complex_rows': first_calculation.template_rows['complex_rows'],
        'object_rows': first_calculation.template_rows['object_rows'],
        'calculation_rows': _combine_word_calculation_rows(calculations),
    }
    return {
        'property_obj': property_group['property_obj'],
        'template_rows': template_rows,
        'calculation_heading': _get_property_group_calculation_heading(
            calculations
        ),
    }


def _combine_word_calculation_rows(calculations):
    """Return calculation rows, adding program labels when needed."""
    if len(calculations) == 1:
        return calculations[0].template_rows['calculation_rows']

    rows = []
    for calculation in calculations:
        rows.append(['Ипотека', calculation.program_label])
        rows.extend(calculation.template_rows['calculation_rows'])
    return rows


def _get_property_group_calculation_heading(calculations):
    """Return the heading for a property mortgage section."""
    headings = {calculation.calculation_heading for calculation in calculations}
    if len(headings) == 1:
        return calculations[0].calculation_heading
    return 'Ипотека'


def _get_word_complex_key(property_obj):
    """Return a stable grouping key for a property complex."""
    if property_obj is None:
        return ('property', None)
    return ('complex', property_obj.building.real_estate_complex_id)


def _get_word_property_key(property_obj):
    """Return a stable grouping key for a property object."""
    if property_obj is None:
        return ('property', None)
    return ('property', property_obj.pk)


def _append_template_table(document, table_element):
    """Append a cloned template table to the document body."""
    cloned_table = deepcopy(table_element)
    _append_body_element(document, cloned_table)
    _append_table_separator_paragraph(document)
    return document.tables[-1]


def _append_table_separator_paragraph(document):
    """Append an empty paragraph so Word keeps adjacent tables separate."""
    _append_body_element(document, OxmlElement('w:p'))


def _append_body_element(document, element):
    """Append a body XML element before section properties."""
    body = document._element.body
    section_properties = None
    for child in body:
        if child.tag == qn('w:sectPr'):
            section_properties = child
            break
    if section_properties is None:
        body.append(element)
    else:
        section_properties.addprevious(element)


def _append_page_break(document):
    """Append a page break to the document."""
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)


def _create_template_document():
    """Создает пустой документ на основе приложенного Word-шаблона."""
    document = Document(str(WORD_TEMPLATE_PATH))
    _apply_document_defaults(document)
    return document


def _apply_document_defaults(document):
    """Настраивает базовый шрифт, не меняя геометрию шаблона."""
    normal_style = document.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(10)


def _populate_template_document(
    document,
    property_obj,
    template_rows,
    calculation_heading,
):
    """Заполняет приложенный Word-шаблон данными расчета."""
    if len(document.tables) < 3:
        _populate_fallback_document(
            document,
            template_rows,
            calculation_heading,
        )
        return

    _populate_complex_table(
        document.tables[0],
        property_obj,
        template_rows['complex_rows'],
    )
    _populate_object_and_calculation_table(
        document.tables[1],
        property_obj,
        template_rows,
        calculation_heading,
    )
    _populate_image_table(document.tables[2], property_obj)


def _populate_fallback_document(
    document,
    template_rows,
    calculation_heading,
):
    """Заполняет документ табличным отчетом, если структура шаблона сломана."""
    _clear_document_body(document)
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(calculation_heading)
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(31, 78, 121)
    _append_key_value_section(
        document,
        'Данные ЖК',
        template_rows['complex_rows'],
    )
    _append_key_value_section(
        document,
        'Объект недвижимости',
        template_rows['object_rows'],
    )
    _append_key_value_section(
        document,
        calculation_heading,
        template_rows['calculation_rows'],
    )


def _clear_document_body(document):
    """Удаляет содержимое тела документа, сохраняя параметры секции."""
    body = document._element.body
    for child in list(body):
        if child.tag != qn('w:sectPr'):
            body.remove(child)


def _populate_complex_table(table, property_obj, fallback_complex_rows):
    """Заполняет верхний блок шаблона данными ЖК."""
    complex_values = _build_complex_template_values(
        property_obj,
        fallback_complex_rows,
    )
    for row_number, label in enumerate(
        complex_values.keys()
    ):
        if row_number >= len(table.rows):
            break
        _set_cell_plain_text(table.rows[row_number].cells[0], label)
        if label == 'Ссылка на карту ЖК':
            _set_cell_hyperlink(
                table.rows[row_number].cells[1],
                complex_values.get(label, ''),
                'Открыть карту ЖК',
            )
        else:
            _set_cell_plain_text(
                table.rows[row_number].cells[1],
                complex_values.get(label, ''),
            )

    if len(table.rows[0].cells) > 2:
        _replace_cell_image(
            table.rows[0].cells[2],
            (
                property_obj.building.real_estate_complex.photo
                if property_obj is not None
                else None
            ),
            Inches(2.25),
        )



def _populate_object_and_calculation_table(
    table,
    property_obj,
    template_rows,
    calculation_heading,
):
    """Заполняет блок объекта и результата расчета в шаблоне."""
    object_values = dict(template_rows['object_rows'])
    calculation_rows = template_rows['calculation_rows']
    _set_cell_plain_text(
        table.rows[0].cells[0],
        'Параметры объекта недвижимости',
    )

    object_labels = [
        'Класс',
        'Сдача',
        'Ключи',
        'Площадь',
        'Этаж',
        'Корпус / квартира',
        'Вид',
        'Отделка',
        'Стоимость',
    ]
    for row_number, label in enumerate(object_labels, start=1):
        if row_number >= len(table.rows):
            break
        _set_cell_plain_text(table.rows[row_number].cells[0], label)
        _set_cell_plain_text(
            table.rows[row_number].cells[1],
            object_values.get(label, ''),
        )

    if len(table.rows) > 10:
        _set_cell_plain_text(table.rows[10].cells[0], calculation_heading)

    _ensure_distinct_value_rows(
        table,
        start_row=11,
        needed_count=len(calculation_rows),
    )
    available_rows = _get_distinct_value_row_numbers(table, start_row=11)
    for row_number, row_data in zip(available_rows, calculation_rows):
        label, value = row_data
        _set_cell_plain_text(table.rows[row_number].cells[0], label)
        _set_cell_plain_text(table.rows[row_number].cells[1], value)

    for row_number in range(11, len(table.rows)):
        if row_number in available_rows[:len(calculation_rows)]:
            continue
        _clear_table_row_text(table.rows[row_number], keep_image_cells=True)

    if len(table.rows[0].cells) > 2:
        _replace_cell_image(
            table.rows[1].cells[2],
            property_obj.layout_image if property_obj is not None else None,
            Inches(2.35),
        )


def _populate_image_table(table, property_obj):
    """Заполняет нижний блок изображений объекта."""
    if len(table.rows) < 2:
        return

    _replace_cell_image(
        table.rows[1].cells[0],
        property_obj.floor_plan_image if property_obj is not None else None,
        Inches(3.1),
    )
    _replace_cell_image(
        table.rows[1].cells[1],
        property_obj.window_view_image if property_obj is not None else None,
        Inches(3.1),
    )


def _get_distinct_value_row_numbers(table, start_row):
    """Возвращает строки с отдельными ячейками подписи и значения."""
    row_numbers = []
    for row_number in range(start_row, len(table.rows)):
        cells = table.rows[row_number].cells
        if len(cells) > 1 and cells[0]._tc is not cells[1]._tc:
            row_numbers.append(row_number)
    return row_numbers


def _ensure_distinct_value_rows(table, start_row, needed_count):
    """Добавляет строки, если в шаблоне их недостаточно для расчета."""
    while len(_get_distinct_value_row_numbers(table, start_row)) < needed_count:
        table.add_row()


def _build_complex_template_values(property_obj, fallback_complex_rows=None):
    """Возвращает значения верхнего блока ЖК."""
    template_labels = [
        'Локация',
        'Метро',
        'Доступность',
        'Жилой комплекс',
        'Класс',
        'Описание',
        'Цель покупки',
        'Кому подходит',
        'Сдача',
        'Ключи',
        'Инвестиционный потенциал',
        'Ссылка на карту ЖК',
    ]
    if property_obj is None:
        fallback_values = dict(fallback_complex_rows or [])
        return {
            label: {
                'Локация': fallback_values.get('Район', ''),
                'Жилой комплекс': fallback_values.get('Жилой комплекс', ''),
            }.get(label, '')
            for label in template_labels
        }

    real_estate_complex = property_obj.building.real_estate_complex
    return {
        'Локация': real_estate_complex.district.name,
        'Метро': _get_complex_metro_text(real_estate_complex),
        'Доступность': _get_complex_metro_availability_text(
            real_estate_complex
        ),
        'Жилой комплекс': real_estate_complex.name,
        'Класс': real_estate_complex.real_estate_class.name,
        'Описание': real_estate_complex.description or '',
        'Цель покупки': '',
        'Кому подходит': '',
        'Сдача': _format_report_value(
            'Сдача',
            property_obj.building.get_commissioning_display(),
        ),
        'Ключи': _format_report_value(
            'Ключи',
            property_obj.building.get_key_handover_display(),
        ),
        'Инвестиционный потенциал': (
            real_estate_complex.investment_potential or ''
        ),
        'Ссылка на карту ЖК': (
            real_estate_complex.map_link or ''
        ),
    }


def _get_complex_metro_text(real_estate_complex):
    """Возвращает список станций метро ЖК."""
    availability = list(real_estate_complex.metro_availability.all())
    if not availability:
        return ''
    return ', '.join(item.metro.station for item in availability)


def _get_complex_metro_availability_text(real_estate_complex):
    """Возвращает текст доступности метро для ЖК."""
    availability = list(real_estate_complex.metro_availability.all())
    if not availability:
        return ''
    return ', '.join(
        f'Пешая, {item.walking_time_minutes} мин.'
        for item in availability
    )


def _build_market_template_rows(report_data):
    """Формирует значения шаблона для обычной ипотеки."""
    property_rows = _build_property_rows(report_data)
    complex_rows, object_rows = _split_complex_and_object_rows(property_rows)
    object_values = _build_object_template_values(
        report_data.property_obj,
        object_rows,
        report_data.final_property_cost,
    )
    result = report_data.result
    calculation_rows = [
        [
            'Первоначальный взнос',
            _format_ruble_value(
                report_data.final_property_cost
                * report_data.initial_payment_percent
                / 100
            ),
        ],
    ]
    calculation_rows.extend(_build_market_monthly_payment_rows(report_data))
    return {
        'complex_rows': complex_rows,
        'object_rows': list(object_values.items()),
        'calculation_rows': calculation_rows,
    }


def _build_trench_template_rows(calculation):
    """Формирует значения шаблона для траншевой ипотеки."""
    _, object_rows = _build_trench_property_sections(calculation)
    object_values = _build_object_template_values(
        calculation.get('property_obj'),
        object_rows,
        calculation['final_property_cost'],
    )
    calculation_rows = [
        ['Первоначальный взнос', _format_ruble_value(calculation['initial_payment'])],
    ]
    calculation_rows.extend(
        _build_trench_monthly_payment_rows(calculation)
    )
    return {
        'complex_rows': [],
        'object_rows': list(object_values.items()),
        'calculation_rows': calculation_rows,
    }


def _build_object_template_values(property_obj, object_rows, property_cost):
    """Формирует значения блока объекта под строки шаблона."""
    values = dict(object_rows)
    template_values = {
        'Класс': '',
        'Сдача': '',
        'Ключи': '',
        'Площадь': values.get('Площадь, м2') or values.get('Площадь') or '',
        'Этаж': values.get('Этаж', ''),
        'Корпус / квартира': '',
        'Вид': '',
        'Отделка': values.get('Отделка', ''),
        'Стоимость': _format_ruble_value(property_cost),
    }

    if property_obj is None:
        return template_values

    real_estate_complex = property_obj.building.real_estate_complex
    template_values.update(
        {
            'Класс': real_estate_complex.real_estate_class.name,
            'Сдача': _format_report_value(
                'Сдача',
                property_obj.building.get_commissioning_display(),
            ),
            'Ключи': _format_report_value(
                'Ключи',
                property_obj.building.get_key_handover_display(),
            ),
            'Площадь': format_compact_decimal(property_obj.area),
            'Этаж': str(property_obj.floor),
            'Корпус / квартира': (
                f'{property_obj.building.number} / '
                f'{property_obj.apartment_number}'
            ),
            'Вид': _get_property_window_view_text(property_obj),
            'Отделка': property_obj.decoration.name,
        }
    )
    return template_values


def _get_property_window_view_text(property_obj):
    """Возвращает виды из окна объекта."""
    window_views = list(property_obj.window_views.all())
    if not window_views:
        return ''
    return ', '.join(window_view.name for window_view in window_views)


def _format_ruble_value(value):
    """Форматирует рублевое значение для шаблонного отчета."""
    if value in (None, ''):
        return ''
    formatted_value = format_currency(value)
    if formatted_value.endswith(',00'):
        formatted_value = formatted_value[:-3]
    return f'{formatted_value} ₽'


def _build_monthly_payment_row(payment_count, end_date, payment_amount):
    """Формирует строку ежемесячного платежа с периодом."""
    month_count = int(payment_count or 0)
    return [
        (
            'Ежемесячный платеж '
            f'на {_format_month_count(month_count)} '
            f'(до {_format_report_value("Дата", end_date)})'
        ),
        _format_ruble_value(payment_amount),
    ]


def _format_month_count(month_count):
    """Возвращает количество месяцев с корректным падежом."""
    remainder_100 = month_count % 100
    remainder_10 = month_count % 10
    if 11 <= remainder_100 <= 14:
        label = 'месяцев'
    elif remainder_10 == 1:
        label = 'месяц'
    elif 2 <= remainder_10 <= 4:
        label = 'месяца'
    else:
        label = 'месяцев'
    return f'{month_count} {label}'


def _build_market_monthly_payment_rows(report_data):
    """Формирует строки ежемесячных платежей рыночной ипотеки."""
    result = report_data.result
    payment_schedule = report_data.payment_schedule or []
    rows = []
    grace_payment_count = int(result.get('grace_payments_count') or 0)
    main_payment_count = int(result.get('main_payments_count') or 0)

    if grace_payment_count:
        rows.append(
            _build_monthly_payment_row(
                grace_payment_count,
                _get_schedule_period_end_date(
                    payment_schedule,
                    0,
                    grace_payment_count,
                    result['grace_period_end_date'],
                ),
                result['grace_monthly_payment'],
            )
        )

    rows.append(
        _build_monthly_payment_row(
            main_payment_count,
            _get_schedule_period_end_date(
                payment_schedule,
                grace_payment_count,
                grace_payment_count + main_payment_count,
                result['mortgage_end_date'],
            ),
            result['main_monthly_payment'],
        )
    )
    return rows


def _get_schedule_period_end_date(
    payment_schedule,
    start_index,
    end_index,
    fallback_date,
):
    """Возвращает дату последнего платежа периода из графика."""
    period_payments = payment_schedule[start_index:end_index]
    if period_payments:
        return period_payments[-1]['payment_date']
    return fallback_date


def _build_trench_monthly_payment_rows(calculation):
    """Формирует строки ежемесячных платежей по траншам."""
    trenches = sorted(calculation['trenches'], key=lambda item: item['date'])
    payment_schedule = calculation.get('payment_schedule') or []
    rows = []

    for trench_index, trench in enumerate(trenches):
        next_trench_date = (
            trenches[trench_index + 1]['date']
            if trench_index + 1 < len(trenches)
            else calculation['mortgage_end_date']
        )
        period_payments = [
            payment
            for payment in payment_schedule
            if trench['date'] <= payment['payment_date'] < next_trench_date
        ]
        payment_count = len(period_payments) or trench.get(
            'payments_count',
            0,
        )
        if period_payments:
            end_date = period_payments[-1]['payment_date']
        else:
            end_date = _calculate_last_payment_date(
                trench['date'],
                payment_count,
            )
        rows.append(
            _build_monthly_payment_row(
                payment_count,
                end_date,
                trench['monthly_payment'],
            )
        )
    return rows


def _calculate_last_payment_date(start_date, payment_count):
    """Возвращает дату последнего платежа периода без графика платежей."""
    if not payment_count:
        return start_date

    from dateutil.relativedelta import relativedelta

    return start_date + relativedelta(months=payment_count - 1)


def _set_cell_plain_text(cell, text):
    """Заменяет текст ячейки, сохраняя табличную разметку."""
    if not cell.paragraphs:
        paragraph = cell.add_paragraph()
    else:
        paragraph = cell.paragraphs[0]

    _clear_paragraph_text_preserving_first_run(paragraph, str(text or ''))
    paragraph.paragraph_format.space_after = Pt(0)
    for extra_paragraph in cell.paragraphs[1:]:
        _clear_paragraph_text_preserving_first_run(extra_paragraph, '')


def _clear_paragraph_text_preserving_first_run(paragraph, text):
    """Заменяет текст параграфа без сброса форматирования первого run."""
    if paragraph.runs:
        first_run = paragraph.runs[0]
    else:
        first_run = paragraph.add_run()

    first_run.text = text
    for run in paragraph.runs[1:]:
        run.text = ''


def _set_cell_hyperlink(cell, url, display_text):
    """Заменяет значение ячейки внешней ссылкой."""
    normalized_url = _normalize_url(url)
    if not normalized_url:
        _set_cell_plain_text(cell, '')
        return

    _clear_cell_content(cell)
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    relationship_id = cell.part.relate_to(
        normalized_url,
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), relationship_id)

    run_element = OxmlElement('w:r')
    run_properties = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    run_properties.append(color)
    run_properties.append(underline)
    run_element.append(run_properties)
    text = OxmlElement('w:t')
    text.text = display_text
    run_element.append(text)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)


def _normalize_url(url):
    """Возвращает URL с протоколом для внешней ссылки."""
    url = str(url or '').strip()
    if not url:
        return ''
    if url.startswith(('http://', 'https://')):
        return url
    return f'https://{url}'


def _clear_table_row_text(row, keep_image_cells=False):
    """Очищает текстовые ячейки строки."""
    for cell in row.cells:
        if keep_image_cells and _cell_has_drawing(cell):
            continue
        _set_cell_plain_text(cell, '')


def _replace_cell_image(cell, image_field, width):
    """Заменяет изображение в ячейке или очищает ее."""
    _clear_cell_content(cell)
    image_path = _get_image_path(image_field)
    if image_path is None:
        cell.add_paragraph()
        return

    paragraph = cell.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=width)


def _get_image_path(image_field):
    """Возвращает путь к изображению Django ImageField."""
    if not image_field:
        return None

    try:
        image_path = Path(image_field.path)
    except (NotImplementedError, ValueError):
        return None

    if image_path.exists():
        return image_path
    return None


def _clear_cell_content(cell):
    """Удаляет все содержимое ячейки."""
    for child in list(cell._tc):
        if child.tag != qn('w:tcPr'):
            cell._tc.remove(child)


def _cell_has_drawing(cell):
    """Проверяет, есть ли в ячейке изображение."""
    return bool(cell._tc.xpath('.//w:drawing'))


def _split_complex_and_object_rows(rows):
    """Разделяет строки на данные ЖК и данные объекта."""
    complex_rows = []
    object_rows = []
    for label, value in rows:
        target_rows = (
            complex_rows if label in COMPLEX_ROW_LABELS else object_rows
        )
        target_rows.append([label, value])
    return complex_rows, object_rows


def _append_key_value_section(document, title, rows):
    """Добавляет раздел с таблицей ключ-значение."""
    if not rows:
        return

    _append_section_heading(document, title)
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    table.autofit = False

    for label, value in rows:
        cells = table.add_row().cells
        cells[0].width = Inches(2.35)
        cells[1].width = Inches(4.95)
        _set_cell_text(cells[0], str(label), bold=True)
        _set_cell_text(cells[1], _format_report_value(label, value))

    _append_spacing(document)


def _append_section_heading(document, title):
    """Добавляет заголовок раздела."""
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(31, 78, 121)


def _append_spacing(document):
    """Добавляет небольшой интервал после таблицы."""
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)


def _set_cell_text(cell, text, bold=False, centered=False):
    """Записывает текст в ячейку с базовым форматированием."""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
        if centered
        else WD_ALIGN_PARAGRAPH.LEFT
    )
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = 'Arial'
    run.font.size = Pt(9)


def _format_report_value(label, value):
    """Форматирует значение для вывода в Word-отчет."""
    if value is None:
        return ''
    if isinstance(value, (date, datetime)):
        return value.strftime('%d.%m.%Y')
    if isinstance(value, str):
        return value
    if 'руб' in label:
        return format_currency(value)
    if '%' in label or label in {'Площадь', 'Площадь, м2'}:
        return format_compact_decimal(value)
    return str(value)


def _build_trench_property_sections(calculation):
    """Формирует строки данных ЖК и объекта для траншевого отчета."""
    property_obj = calculation.get('property_obj')
    if property_obj is None:
        return [], _build_trench_price_rows(calculation)

    real_estate_complex = property_obj.building.real_estate_complex
    complex_rows = [
        ['Город', real_estate_complex.district.city.name],
        ['Район', real_estate_complex.district.name],
        ['Застройщик', real_estate_complex.developer.name],
        ['Жилой комплекс', real_estate_complex.name],
        ['Класс ЖК', real_estate_complex.real_estate_class.name],
    ]
    commissioning = property_obj.building.get_commissioning_display()
    key_handover = property_obj.building.get_key_handover_display()
    if commissioning:
        complex_rows.append(['Сдача', commissioning])
    if key_handover:
        complex_rows.append(['Ключи', key_handover])

    object_rows = [
        ['Корпус', property_obj.building.number],
        ['Номер квартиры', property_obj.apartment_number],
        ['Площадь, м2', property_obj.area],
        ['Планировка', property_obj.layout.name],
        ['Этаж', property_obj.floor],
        ['Отделка', property_obj.decoration.name],
    ]
    object_rows.extend(_build_trench_price_rows(calculation))
    return complex_rows, object_rows


def _build_trench_price_rows(calculation):
    """Формирует строки стоимостных показателей траншевого отчета."""
    percent_label, rubles_label = get_discount_markup_labels(
        calculation['discount_markup_type']
    )
    return [
        ['Базовая стоимость объекта, руб.', calculation['base_property_cost']],
        [percent_label, calculation['discount_markup_value']],
        [
            rubles_label,
            _calculate_trench_discount_markup_amount(calculation),
        ],
        [
            'Итоговая стоимость объекта, руб.',
            calculation['final_property_cost'],
        ],
    ]


def _calculate_trench_discount_markup_amount(calculation):
    """Возвращает сумму скидки или удорожания для траншевого отчета."""
    base_property_cost = float(calculation['base_property_cost'])
    discount_markup_value = float(calculation['discount_markup_value'])
    return base_property_cost * discount_markup_value / 100
